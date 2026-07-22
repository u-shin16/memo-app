from __future__ import annotations

import json
import base64
import unittest
from io import BytesIO
from unittest.mock import Mock, patch

from docx import Document
from openpyxl import Workbook
from pypdf import PdfWriter

import app as app_module


ONE_PIXEL_PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
)


class AiFileExtractionTests(unittest.TestCase):
    def test_extracts_utf8_and_cp932_text(self):
        self.assertEqual(
            app_module.extract_ai_file_text("sample.txt", "テスト本文".encode()),
            "テスト本文",
        )
        self.assertEqual(
            app_module.extract_ai_file_text("sample.csv", "項目,値\n件数,3".encode("cp932")),
            "項目,値\n件数,3",
        )

    def test_formats_json(self):
        text = app_module.extract_ai_file_text("sample.json", b'{"topic":"memo"}')
        self.assertEqual(json.loads(text), {"topic": "memo"})

    def test_extracts_docx_paragraphs_and_tables(self):
        document = Document()
        document.add_paragraph("見出し")
        table = document.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "項目"
        table.cell(0, 1).text = "内容"
        stream = BytesIO()
        document.save(stream)

        text = app_module.extract_ai_file_text("sample.docx", stream.getvalue())

        self.assertIn("見出し", text)
        self.assertIn("項目\t内容", text)

    def test_extracts_xlsx_sheets_and_cells(self):
        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "計画"
        sheet.append(["作業", "期限"])
        sheet.append(["実装", "6月"])
        stream = BytesIO()
        workbook.save(stream)

        text = app_module.extract_ai_file_text("sample.xlsx", stream.getvalue())

        self.assertIn("[シート: 計画]", text)
        self.assertIn("実装\t6月", text)

    def test_rejects_blank_pdf_and_unsupported_file(self):
        writer = PdfWriter()
        writer.add_blank_page(width=100, height=100)
        stream = BytesIO()
        writer.write(stream)

        with self.assertRaisesRegex(ValueError, "文字を抽出できません"):
            app_module.extract_ai_file_text("blank.pdf", stream.getvalue())
        with self.assertRaisesRegex(ValueError, "対応していない"):
            app_module.extract_ai_file_text("sample.pptx", b"data")

    def test_rejects_oversized_and_broken_office_files(self):
        with patch.object(app_module, "_AI_FILE_MAX_BYTES", 3):
            with self.assertRaisesRegex(ValueError, "10MB以下"):
                app_module.extract_ai_file_text("large.txt", b"1234")
        with self.assertRaisesRegex(ValueError, "破損"):
            app_module.extract_ai_file_text("broken.docx", b"not-a-zip")

    def test_validates_supported_image_signatures(self):
        self.assertEqual(
            app_module._get_validated_image_mime_type("photo.png", ONE_PIXEL_PNG),
            "image/png",
        )
        heic_header = b"\x00\x00\x00\x18ftypheic" + (b"\x00" * 12)
        self.assertEqual(
            app_module._get_validated_image_mime_type("photo.heic", heic_header),
            "image/heic",
        )
        with self.assertRaisesRegex(ValueError, "拡張子と一致"):
            app_module._get_validated_image_mime_type("fake.jpg", ONE_PIXEL_PNG)


class AiFileApiTests(unittest.TestCase):
    def setUp(self):
        app_module.app.config.update(TESTING=True)
        self.client = app_module.app.test_client()

    @staticmethod
    def _mock_client(tree):
        client = Mock()
        client.models.generate_content.return_value.text = json.dumps(tree, ensure_ascii=False)
        return client

    def test_note_api_sends_prompt_and_file_text_to_model(self):
        client = self._mock_client({"title": "結果", "content": "本文", "children": []})
        with patch.object(app_module, "create_gemini_client", return_value=client):
            response = self.client.post(
                "/api/ai-note",
                data={
                    "prompt": "要点をまとめて",
                    "file": (BytesIO("参照資料の本文".encode()), "reference.txt"),
                },
                content_type="multipart/form-data",
            )

        self.assertEqual(response.status_code, 200)
        contents = client.models.generate_content.call_args.kwargs["contents"]
        self.assertIn("要点をまとめて", contents)
        self.assertIn("reference.txt", contents)
        self.assertIn("参照資料の本文", contents)

    def test_mindmap_api_keeps_json_request_compatibility(self):
        client = self._mock_client({"title": "マップ", "memo": "", "children": []})
        with patch.object(app_module, "create_gemini_client", return_value=client):
            response = self.client.post("/api/ai-mindmap", json={"prompt": "学習計画"})

        self.assertEqual(response.status_code, 200)
        contents = client.models.generate_content.call_args.kwargs["contents"]
        self.assertIn("学習計画", contents)
        self.assertIn("すべてのノードで memo を必ず生成", contents)

    def test_mindmap_api_normalizes_descriptions_to_memo(self):
        client = self._mock_client({
            "title": "計画",
            "content": "全体の進め方を整理する",
            "children": [
                {
                    "title": "準備",
                    "description": "必要な資料と担当を決める",
                    "children": [
                        {"title": "調査", "summary": "先行事例を確認する", "children": []},
                    ],
                },
            ],
        })
        with patch.object(app_module, "create_gemini_client", return_value=client):
            response = self.client.post("/api/ai-mindmap", json={"prompt": "学習計画"})

        self.assertEqual(response.status_code, 200)
        tree = response.get_json()["tree"]
        self.assertEqual(tree["memo"], "全体の進め方を整理する")
        self.assertEqual(tree["children"][0]["memo"], "必要な資料と担当を決める")
        self.assertEqual(tree["children"][0]["children"][0]["memo"], "先行事例を確認する")

    def test_mindmap_api_fills_missing_memo(self):
        client = self._mock_client({"title": "計画", "memo": "", "children": [{"title": "準備"}]})
        with patch.object(app_module, "create_gemini_client", return_value=client):
            response = self.client.post("/api/ai-mindmap", json={"prompt": "学習計画"})

        self.assertEqual(response.status_code, 200)
        tree = response.get_json()["tree"]
        self.assertTrue(tree["memo"].strip())
        self.assertTrue(tree["children"][0]["memo"].strip())

    def test_accepts_file_without_prompt(self):
        client = self._mock_client({"title": "結果", "content": "本文", "children": []})
        with patch.object(app_module, "create_gemini_client", return_value=client):
            response = self.client.post(
                "/api/ai-note",
                data={"file": (BytesIO("ファイルだけで生成".encode()), "only.md")},
                content_type="multipart/form-data",
            )

        self.assertEqual(response.status_code, 200)
        contents = client.models.generate_content.call_args.kwargs["contents"]
        self.assertIn("添付ファイルの内容を整理してください", contents)

    def test_note_api_sends_image_as_multimodal_input(self):
        client = self._mock_client({"title": "画像メモ", "content": "説明", "children": []})
        with patch.object(app_module, "create_gemini_client", return_value=client):
            response = self.client.post(
                "/api/ai-note",
                data={
                    "prompt": "画像を説明して",
                    "file": (BytesIO(ONE_PIXEL_PNG), "sample.png"),
                },
                content_type="multipart/form-data",
            )

        self.assertEqual(response.status_code, 200)
        contents = client.models.generate_content.call_args.kwargs["contents"]
        self.assertIsInstance(contents, list)
        self.assertEqual(contents[0].inline_data.mime_type, "image/png")
        self.assertEqual(contents[0].inline_data.data, ONE_PIXEL_PNG)
        self.assertIn("画像を説明して", contents[1])
        self.assertIn("sample.png", contents[1])

    def test_mindmap_api_accepts_image_without_prompt(self):
        client = self._mock_client({"title": "画像マップ", "memo": "", "children": []})
        with patch.object(app_module, "create_gemini_client", return_value=client):
            response = self.client.post(
                "/api/ai-mindmap",
                data={"file": (BytesIO(ONE_PIXEL_PNG), "only.png")},
                content_type="multipart/form-data",
            )

        self.assertEqual(response.status_code, 200)
        contents = client.models.generate_content.call_args.kwargs["contents"]
        self.assertIn("添付画像の内容を読み取り", contents[1])

    def test_rejects_image_with_mismatched_content(self):
        with patch.object(app_module, "create_gemini_client") as create_client:
            response = self.client.post(
                "/api/ai-note",
                data={"file": (BytesIO(b"not-an-image"), "fake.png")},
                content_type="multipart/form-data",
            )

        self.assertEqual(response.status_code, 400)
        self.assertIn("拡張子と一致", response.get_json()["error"])
        create_client.assert_not_called()

    def test_rejects_request_without_prompt_or_file(self):
        with patch.object(app_module, "create_gemini_client") as create_client:
            response = self.client.post("/api/ai-note", json={"prompt": ""})

        self.assertEqual(response.status_code, 400)
        self.assertIn("ファイルを選択", response.get_json()["error"])
        create_client.assert_not_called()

    def test_index_exposes_supported_image_types(self):
        response = self.client.get("/app")
        html = response.get_data(as_text=True)

        self.assertEqual(response.status_code, 200)
        self.assertEqual(html.count(".png,.jpg,.jpeg,.webp,.heic,.heif"), 2)
        self.assertIn('id="mindMapDeleteBtn"', html)
        self.assertIn('class="settings-action-btn danger"', html)
        self.assertIn('id="appManageBtn"', html)
        self.assertIn('id="mindMapManageBtn"', html)
        self.assertIn('id="appManagementOverlay"', html)
        self.assertIn('id="appManagementTitle" class="templates-title">アプリ設定</p>', html)
        self.assertEqual(html.count('aria-label="アプリ設定を開く"'), 2)
        self.assertIn('id="appHowToBtn"', html)
        self.assertIn('id="appHowToDialog"', html)
        self.assertIn('id="appDataInfoBtn"', html)
        self.assertIn('id="appAccountBtn"', html)
        self.assertIn('aria-controls="appAccountDialog"', html)
        self.assertIn('id="appAccountDialog"', html)
        self.assertEqual(html.count('aria-controls="appAccountDialog"'), 3)
        self.assertIn('id="appAccountEditDisplayNameBtn"', html)
        self.assertIn('id="appAccountLogoutBtn"', html)
        self.assertIn('id="appAccountDeleteBtn"', html)
        self.assertIn('id="authGoogleBtn"', html)
        self.assertIn("Googleで続ける", html)
        self.assertIn('id="googleLinkOverlay"', html)
        self.assertIn('id="appCreatorInfoBtn"', html)
        self.assertIn('id="appContactBtn"', html)
        self.assertIn('id="appInfoDialog"', html)
        self.assertIn('id="appNoteHowToTab"', html)
        self.assertIn('id="appMindMapHowToTab"', html)
        self.assertIn('id="appNoteHowToPanel"', html)
        self.assertIn('id="appMindMapHowToPanel"', html)
        self.assertIn("メモを作成・編集する", html)
        self.assertIn("マインドマップを作成・切り替える", html)
        self.assertIn("テンプレートと「戻す」を使う", html)
        self.assertIn("制作者情報", html)
        self.assertIn('href="https://profile.webtool-labs.com/"', html)
        self.assertIn("お問い合わせフォーム", html)
        self.assertIn('href="https://docs.google.com/forms/d/e/1FAIpQLSePfOxSZwYoGcL7csdt1RbLY4eQ9gdu6ePqcgZ96xTxZj8GXA/viewform?usp=publish-editor"', html)
        self.assertNotIn("github.com/u-shin16/memo-app", html)


if __name__ == "__main__":
    unittest.main()
