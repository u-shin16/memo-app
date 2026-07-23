from __future__ import annotations

import json
import os
import urllib.error
import urllib.request
import zipfile
from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path

from google import genai
from google.genai import types as genai_types
from dotenv import load_dotenv
from flask import Flask, Response, abort, jsonify, make_response, render_template, request
from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader

load_dotenv()

BASE_DIR = Path(__file__).resolve().parent
TEMPLATES_DIR = BASE_DIR / "templates"
app = Flask(__name__)

# ── SEOページ関連の設定 ────────────────────────────────────────────────────────
# アプリ本体（/app）とは切り離した、検索エンジン向けの静的ページ群で使う定数・
# ページ一覧。sitemap.xml もこの一覧から自動生成する。

SITE_NAME = "まとめときや"
SITE_URL = "https://matome.webtool-labs.com"
FIREBASE_AUTH_DOMAIN = "memo-app-9dd98.firebaseapp.com"
OPERATOR_NAME = "WebTool-Labs"
OPERATOR_PROFILE_URL = "https://profile.webtool-labs.com/"
OPERATOR_SITE_URL = "https://webtool-labs.com/"
CONTACT_FORM_URL = (
    "https://docs.google.com/forms/d/e/1FAIpQLSePfOxSZwYoGcL7csdt1RbLY4eQ9gdu6ePqcgZ96xTxZj8GXA/viewform?usp=publish-editor"
)

# 狙うキーワードごとの1ページ。テンプレートは templates/seo/keywords/ 配下。
KEYWORD_PAGES = [
    {
        "slug": "hierarchical-memo",
        "keyword": "階層メモ アプリ",
        "label": "階層メモアプリとは",
        "template": "seo/keywords/hierarchical-memo.html",
    },
    {
        "slug": "free-mindmap",
        "keyword": "マインドマップ 作成 無料",
        "label": "無料マインドマップ作成",
        "template": "seo/keywords/free-mindmap.html",
    },
    {
        "slug": "organize-ideas",
        "keyword": "メモ アイデア整理",
        "label": "アイデア整理のコツ",
        "template": "seo/keywords/organize-ideas.html",
    },
    {
        "slug": "ai-memo-generator",
        "keyword": "AI メモ 自動生成",
        "label": "AIでメモを自動生成",
        "template": "seo/keywords/ai-memo-generator.html",
    },
    {
        "slug": "collaborative-memo",
        "keyword": "共同編集 メモ帳",
        "label": "共同編集メモ帳",
        "template": "seo/keywords/collaborative-memo.html",
    },
    {
        "slug": "structured-notes",
        "keyword": "ノート 構造化",
        "label": "ノートを構造化する",
        "template": "seo/keywords/structured-notes.html",
    },
]
KEYWORD_PAGES_BY_SLUG = {page["slug"]: page for page in KEYWORD_PAGES}

# 悩み系キーワードで書くブログ記事。テンプレートは templates/seo/blog/ 配下。
BLOG_POSTS = [
    {
        "slug": "organize-scattered-ideas",
        "title": "アイデアが頭の中でごちゃごちゃになる人へ。階層メモで思考を整理するコツ",
        "template": "seo/blog/organize-scattered-ideas.html",
    },
    {
        "slug": "mindmap-vs-memo",
        "title": "マインドマップとメモ、結局どっちを使えばいい？使い分けの考え方",
        "template": "seo/blog/mindmap-vs-memo.html",
    },
    {
        "slug": "ai-brainstorming-tips",
        "title": "一人ブレストで手が止まったときの、AIとメモアプリの組み合わせ方",
        "template": "seo/blog/ai-brainstorming-tips.html",
    },
]
BLOG_POSTS_BY_SLUG = {post["slug"]: post for post in BLOG_POSTS}

# /faq とランディングページ抜粋、FAQPage構造化データで共用する質問一覧。
FAQ_ITEMS = [
    {
        "question": "無料で使えますか？",
        "answer": "はい、まとめときやは無料でご利用いただけます。メモ作成・階層管理・マインドマップ・AI生成・共同編集など、主要な機能はすべて追加費用なしで使えます。",
    },
    {
        "question": "会員登録は必要ですか？",
        "answer": "まとめときやを使うにはアカウント登録（メールアドレスまたはGoogleアカウント）が必要です。登録すればログインしたどの端末からでも続きから編集できます。",
    },
    {
        "question": "作成したメモはどこに保存されますか？",
        "answer": "メモはお使いのアカウントに紐づけてクラウド上に保存されます。端末を変えても、ログインすれば同じ内容を確認・編集できます。",
    },
    {
        "question": "他の人と一緒に編集できますか？",
        "answer": "はい。「共同編集」機能を使うと、合言葉を伝えるだけで複数人が同じメモをリアルタイムで編集できます。",
    },
    {
        "question": "AIでメモを作るとき、入力した内容は保存されますか？",
        "answer": "AI生成のために入力したテーマや添付ファイルは、生成処理のためだけに利用され、アップロード内容自体をサーバー側に保存することはありません。",
    },
    {
        "question": "作成したメモを他の形式で書き出せますか？",
        "answer": "メモはPDF・テキスト・Markdown形式で、マインドマップはPNG・SVG・PDF形式でダウンロードできます。",
    },
]


def get_public_pages() -> list[dict]:
    """sitemap.xml生成用の公開ページ一覧（パス・優先度・対応テンプレート）。
    アプリ内API・ログイン後にしか意味を持たないページは含めない。"""
    pages = [
        {"path": "/", "template": "seo/landing.html", "priority": "1.0"},
        {"path": "/app", "template": "index.html", "priority": "0.9"},
        {"path": "/how-to-use", "template": "seo/how_to_use.html", "priority": "0.6"},
        {"path": "/faq", "template": "seo/faq.html", "priority": "0.6"},
        {"path": "/about", "template": "seo/about.html", "priority": "0.4"},
        {"path": "/privacy", "template": "seo/privacy.html", "priority": "0.3"},
        {"path": "/terms", "template": "seo/terms.html", "priority": "0.3"},
        {"path": "/contact", "template": "seo/contact.html", "priority": "0.4"},
    ]
    for page in KEYWORD_PAGES:
        pages.append({"path": f"/{page['slug']}", "template": page["template"], "priority": "0.8"})
    for post in BLOG_POSTS:
        pages.append({"path": f"/blog/{post['slug']}", "template": post["template"], "priority": "0.6"})
    return pages


@app.context_processor
def inject_seo_globals():
    def static_version(filename: str) -> int:
        try:
            return int((BASE_DIR / "static" / filename).stat().st_mtime)
        except OSError:
            return 0

    return {
        "site_name": SITE_NAME,
        "site_url": SITE_URL,
        "operator_name": OPERATOR_NAME,
        "operator_profile_url": OPERATOR_PROFILE_URL,
        "operator_site_url": OPERATOR_SITE_URL,
        "contact_form_url": CONTACT_FORM_URL,
        "keyword_pages": KEYWORD_PAGES,
        "blog_posts": BLOG_POSTS,
        "canonical_url": f"{SITE_URL}{request.path}",
        "current_year": datetime.now(tz=timezone.utc).year,
        "static_version": static_version,
    }

_USE_VERTEX_AI = os.getenv("USE_VERTEX_AI", "false").lower() == "true"
_GEMINI_MODEL = os.getenv("GEMINI_MODEL", "gemini-2.5-flash-lite")
_AI_FILE_MAX_BYTES = 10 * 1024 * 1024
_AI_FILE_TEXT_MAX_CHARS = 100_000
_AI_OFFICE_MAX_UNCOMPRESSED_BYTES = 50 * 1024 * 1024
_AI_DOCUMENT_EXTENSIONS = {".txt", ".md", ".csv", ".json", ".pdf", ".docx", ".xlsx"}
_AI_IMAGE_MIME_TYPES = {
    ".png": "image/png",
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".webp": "image/webp",
    ".heic": "image/heic",
    ".heif": "image/heif",
}
_AI_FILE_EXTENSIONS = _AI_DOCUMENT_EXTENSIONS | set(_AI_IMAGE_MIME_TYPES)

app.config["MAX_CONTENT_LENGTH"] = _AI_FILE_MAX_BYTES + (512 * 1024)

_NOTE_PROMPT = """\
あなたは階層型メモ生成AIです。与えられたテーマをもとに、階層構造のメモをJSON形式で生成してください。

出力はJSONオブジェクトのみ（マークダウンコードブロックや説明文は不要）:
{{
  "title": "ルートメモのタイトル",
  "content": "ルートメモの内容（概要や導入文）",
  "children": [
    {{
      "title": "セクション見出し",
      "content": "このセクションの詳細な内容",
      "children": [
        {{
          "title": "サブ項目",
          "content": "サブ項目の内容",
          "children": []
        }}
      ]
    }}
  ]
}}

ルール:
- 入力と同じ言語で出力する
- ルート: テーマ全体のタイトルと概要
- セクション: 3〜5個のメインセクション
- 各セクション: 具体的な説明文を content に記載
- サブ項目: 必要に応じて1〜3個
- content は文章で書く（箇条書き不可）
- 出力: JSONオブジェクトのみ

テーマ：{prompt}"""

_MINDMAP_PROMPT = """\
あなたはマインドマップ生成AIです。与えられたテーマや内容をもとに、マインドマップをJSON形式で生成してください。

出力はJSONオブジェクトのみ（マークダウンコードブロックや説明文は不要）:
{{
  "title": "ルートノードのタイトル",
  "memo": "テーマ全体の要約、背景、重要な観点を1〜3文で説明する",
  "children": [
    {{
      "title": "ブランチ",
      "memo": "このブランチで扱う内容、なぜ重要か、具体的に考えるポイントを1〜3文で説明する",
      "children": [
        {{
          "title": "サブ項目",
          "memo": "このサブ項目の具体的な内容、例、実行のヒントを1〜2文で説明する",
          "children": []
        }}
      ]
    }}
  ]
}}

ルール:
- 入力と同じ言語で出力する（日本語入力→日本語ノード）
- すべてのノードで memo を必ず生成する（空文字は禁止）
- memo はノード本文として読める自然な文章にする
- 詳細説明・背景・具体例・注意点・次のアクションは title ではなく memo に書く
- ルートタイトル：簡潔に（15文字以内推奨）
- メインブランチ：3〜6個
- 各ブランチのサブ項目：1〜4個
- 最大深さ：ルート＋2レベル
- ノードタイトル：短く明確に（20文字以内推奨）
- 出力：JSONオブジェクトのみ

テーマ・内容：{prompt}"""


def create_gemini_client():
    api_key = os.getenv("GEMINI_API_KEY", "").strip()
    if _USE_VERTEX_AI:
        credentials_path = os.getenv("GOOGLE_APPLICATION_CREDENTIALS", "").strip()
        if credentials_path and not Path(credentials_path).expanduser().is_file():
            if api_key:
                return genai.Client(api_key=api_key)
            raise RuntimeError(
                "Vertex AIの認証ファイルが見つかりません。"
                "GEMINI_API_KEYを設定するか、GOOGLE_APPLICATION_CREDENTIALSのパスを確認してください。"
            )
        project = os.getenv("GOOGLE_CLOUD_PROJECT", "")
        location = os.getenv("GOOGLE_CLOUD_LOCATION", "us-central1")
        if not project:
            if api_key:
                return genai.Client(api_key=api_key)
            raise RuntimeError("GOOGLE_CLOUD_PROJECT が設定されていません")
        return genai.Client(vertexai=True, project=project, location=location)

    if not api_key:
        raise RuntimeError("GEMINI_API_KEY が設定されていません")
    return genai.Client(api_key=api_key)


def _decode_text_file(data: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "cp932"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ValueError("ファイルの文字コードを読み取れませんでした。UTF-8形式で保存してください。")


def _validate_office_archive(data: bytes) -> None:
    try:
        with zipfile.ZipFile(BytesIO(data)) as archive:
            total_size = sum(info.file_size for info in archive.infolist())
    except zipfile.BadZipFile as exc:
        raise ValueError("Officeファイルが破損しているか、形式が正しくありません。") from exc
    if total_size > _AI_OFFICE_MAX_UNCOMPRESSED_BYTES:
        raise ValueError("展開後のファイルサイズが大きすぎます。")


def _extract_pdf_text(data: bytes) -> str:
    reader = PdfReader(BytesIO(data))
    if reader.is_encrypted:
        try:
            unlocked = reader.decrypt("")
        except Exception as exc:
            raise ValueError("パスワード付きPDFは読み込めません。") from exc
        if not unlocked:
            raise ValueError("パスワード付きPDFは読み込めません。")
    return "\n\n".join(page.extract_text() or "" for page in reader.pages)


def _extract_docx_text(data: bytes) -> str:
    _validate_office_archive(data)
    document = Document(BytesIO(data))
    chunks = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells]
            if any(values):
                chunks.append("\t".join(values))
    return "\n".join(chunks)


def _extract_xlsx_text(data: bytes) -> str:
    _validate_office_archive(data)
    workbook = load_workbook(BytesIO(data), read_only=True, data_only=False)
    chunks = []
    try:
        for sheet in workbook.worksheets:
            chunks.append(f"[シート: {sheet.title}]")
            for row in sheet.iter_rows(values_only=True):
                values = ["" if value is None else str(value) for value in row]
                if any(value.strip() for value in values):
                    chunks.append("\t".join(values))
    finally:
        workbook.close()
    return "\n".join(chunks)


def _get_validated_image_mime_type(filename: str, data: bytes) -> str:
    extension = Path(filename).suffix.lower()
    mime_type = _AI_IMAGE_MIME_TYPES.get(extension)
    if not mime_type:
        raise ValueError("対応していない画像形式です。")
    if not data:
        raise ValueError("ファイルが空です。")
    if len(data) > _AI_FILE_MAX_BYTES:
        raise ValueError("ファイルサイズは10MB以下にしてください。")

    is_valid = False
    if extension == ".png":
        is_valid = data.startswith(b"\x89PNG\r\n\x1a\n")
    elif extension in {".jpg", ".jpeg"}:
        is_valid = data.startswith(b"\xff\xd8\xff")
    elif extension == ".webp":
        is_valid = len(data) >= 12 and data.startswith(b"RIFF") and data[8:12] == b"WEBP"
    else:
        heif_brands = {b"heic", b"heix", b"hevc", b"hevx", b"mif1", b"msf1"}
        is_valid = (
            len(data) >= 12
            and data[4:8] == b"ftyp"
            and data[8:12] in heif_brands
        )

    if not is_valid:
        raise ValueError("画像ファイルが破損しているか、形式が拡張子と一致しません。")
    return mime_type


def extract_ai_file_text(filename: str, data: bytes) -> str:
    extension = Path(filename).suffix.lower()
    if extension not in _AI_DOCUMENT_EXTENSIONS:
        supported = ", ".join(sorted(_AI_FILE_EXTENSIONS))
        raise ValueError(f"対応していないファイル形式です。対応形式: {supported}")
    if not data:
        raise ValueError("ファイルが空です。")
    if len(data) > _AI_FILE_MAX_BYTES:
        raise ValueError("ファイルサイズは10MB以下にしてください。")

    try:
        if extension in {".txt", ".md", ".csv"}:
            text = _decode_text_file(data)
        elif extension == ".json":
            parsed = json.loads(_decode_text_file(data))
            text = json.dumps(parsed, ensure_ascii=False, indent=2)
        elif extension == ".pdf":
            text = _extract_pdf_text(data)
        elif extension == ".docx":
            text = _extract_docx_text(data)
        else:
            text = _extract_xlsx_text(data)
    except ValueError:
        raise
    except Exception as exc:
        raise ValueError("ファイルの内容を読み取れませんでした。") from exc

    text = text.replace("\x00", "").strip()
    if not text:
        raise ValueError("ファイルから文字を抽出できませんでした。画像のみのPDFには対応していません。")
    if len(text) > _AI_FILE_TEXT_MAX_CHARS:
        text = text[:_AI_FILE_TEXT_MAX_CHARS] + "\n\n[以降は文字数上限のため省略されました]"
    return text


def get_ai_input_from_request() -> tuple[str, genai_types.Part | None]:
    if request.is_json:
        data = request.get_json(silent=True) or {}
        prompt = str(data.get("prompt", "")).strip()
        upload = None
    else:
        prompt = str(request.form.get("prompt", "")).strip()
        upload = request.files.get("file")

    file_text = ""
    filename = ""
    image_part = None
    if upload and upload.filename:
        filename = upload.filename.replace("\\", "/").rsplit("/", 1)[-1][:255]
        file_data = upload.read(_AI_FILE_MAX_BYTES + 1)
        extension = Path(filename).suffix.lower()
        if extension in _AI_IMAGE_MIME_TYPES:
            mime_type = _get_validated_image_mime_type(filename, file_data)
            image_part = genai_types.Part.from_bytes(data=file_data, mime_type=mime_type)
        else:
            file_text = extract_ai_file_text(filename, file_data)

    if not prompt and not file_text and image_part is None:
        raise ValueError("テーマを入力するか、ファイルを選択してください。")
    if image_part is not None:
        instruction = prompt or "添付画像の内容を読み取り、重要な情報を整理してください。"
        return f"依頼・テーマ:\n{instruction}\n\n添付画像: {filename}", image_part
    if not file_text:
        return prompt, None

    instruction = prompt or "添付ファイルの内容を整理してください。"
    return (
        f"依頼・テーマ:\n{instruction}\n\n"
        f"添付ファイル「{filename}」の内容:\n"
        "--- ファイル内容ここから ---\n"
        f"{file_text}\n"
        "--- ファイル内容ここまで ---"
    ), None


def build_ai_contents(prompt_template: str, prompt: str, image_part: genai_types.Part | None):
    formatted_prompt = prompt_template.format(prompt=prompt)
    if image_part is None:
        return formatted_prompt
    return [image_part, formatted_prompt]


def _extract_mindmap_memo(node: dict, title: str, depth: int) -> str:
    for key in ("memo", "content", "description", "summary", "detail", "details"):
        value = node.get(key)
        if isinstance(value, list):
            text = "\n".join(str(item).strip() for item in value if str(item).strip())
        else:
            text = str(value or "").strip()
        if text:
            return text

    if depth == 0:
        return f"{title}の全体像、主要な観点、深掘りすべきポイントを整理します。"
    return f"{title}について、目的・要点・具体例を整理するためのメモです。"


def normalize_ai_mindmap_tree(tree):
    if not isinstance(tree, dict):
        raise ValueError("AIの出力形式が正しくありません。")

    def normalize_node(node, depth=0):
        if not isinstance(node, dict):
            return None
        title = str(node.get("title") or "トピック").strip()[:80] or "トピック"
        children = node.get("children")
        if not isinstance(children, list):
            children = []
        return {
            "title": title,
            "memo": _extract_mindmap_memo(node, title, depth),
            "children": [
                child
                for child in (normalize_node(child_node, depth + 1) for child_node in children)
                if child
            ],
        }

    return normalize_node(tree)


@app.get("/app")
def index():
    response = make_response(render_template("index.html"))
    response.headers["Cache-Control"] = "no-store, max-age=0"
    return response


@app.get("/auth/action")
def auth_action():
    response = make_response(render_template("auth_action.html"))
    response.headers["Cache-Control"] = "no-store, max-age=0"
    return response


def proxy_firebase_helper(path: str) -> Response:
    query = request.query_string.decode("utf-8")
    target_url = f"https://{FIREBASE_AUTH_DOMAIN}/{path}"
    if query:
        target_url = f"{target_url}?{query}"

    headers = {
        key: value
        for key, value in request.headers.items()
        if key.lower() not in {"host", "content-length", "connection", "accept-encoding"}
    }
    data = request.get_data() if request.method not in {"GET", "HEAD"} else None
    upstream_request = urllib.request.Request(
        target_url,
        data=data,
        headers=headers,
        method=request.method,
    )

    try:
        with urllib.request.urlopen(upstream_request, timeout=15) as upstream:
            body = upstream.read()
            response_headers = {
                key: value
                for key, value in upstream.headers.items()
                if key.lower() not in {"connection", "transfer-encoding", "content-encoding", "content-length"}
            }
            return Response(body, status=upstream.status, headers=response_headers)
    except urllib.error.HTTPError as error:
        body = error.read()
        response_headers = {
            key: value
            for key, value in error.headers.items()
            if key.lower() not in {"connection", "transfer-encoding", "content-encoding", "content-length"}
        }
        return Response(body, status=error.code, headers=response_headers)
    except urllib.error.URLError as error:
        return Response(f"Firebase auth helper proxy failed: {error.reason}", status=502)


@app.route("/__/auth/<path:auth_path>", methods=["GET", "POST", "HEAD", "OPTIONS"])
def firebase_auth_helper(auth_path: str):
    return proxy_firebase_helper(f"__/auth/{auth_path}")


@app.route("/__/firebase/init.json", methods=["GET", "HEAD"])
def firebase_init_json():
    return proxy_firebase_helper("__/firebase/init.json")


# ── SEOページ ─────────────────────────────────────────────────────────────────

@app.get("/")
def seo_landing():
    webapp_jsonld = {
        "@context": "https://schema.org",
        "@type": "WebApplication",
        "name": SITE_NAME,
        "url": f"{SITE_URL}/app",
        "description": (
            "頭の中のアイデアやタスクを、階層型のメモとマインドマップで整理できる無料アプリ。"
            "AIによる自動生成や複数人での共同編集にも対応しています。"
        ),
        "applicationCategory": "ProductivityApplication",
        "operatingSystem": "Web",
        "offers": {
            "@type": "Offer",
            "price": "0",
            "priceCurrency": "JPY",
        },
    }
    return render_template("seo/landing.html", webapp_jsonld=webapp_jsonld, faq_items=FAQ_ITEMS[:3])


@app.get("/how-to-use")
def how_to_use():
    return render_template("seo/how_to_use.html")


@app.get("/faq")
def faq():
    faq_jsonld = {
        "@context": "https://schema.org",
        "@type": "FAQPage",
        "mainEntity": [
            {
                "@type": "Question",
                "name": item["question"],
                "acceptedAnswer": {"@type": "Answer", "text": item["answer"]},
            }
            for item in FAQ_ITEMS
        ],
    }
    return render_template("seo/faq.html", faq_items=FAQ_ITEMS, faq_jsonld=faq_jsonld)


@app.get("/about")
def about():
    return render_template("seo/about.html")


@app.get("/privacy")
def privacy():
    return render_template("seo/privacy.html")


@app.get("/terms")
def terms():
    return render_template("seo/terms.html")


@app.get("/contact")
def contact():
    return render_template("seo/contact.html")


@app.get("/robots.txt")
def robots_txt():
    body = "\n".join([
        "User-agent: *",
        "Allow: /",
        f"Sitemap: {SITE_URL}/sitemap.xml",
        "",
    ])
    return Response(body, mimetype="text/plain")


@app.get("/sitemap.xml")
def sitemap_xml():
    entries = []
    for page in get_public_pages():
        template_path = TEMPLATES_DIR / page["template"]
        try:
            mtime = template_path.stat().st_mtime
        except OSError:
            mtime = datetime.now(tz=timezone.utc).timestamp()
        lastmod = datetime.fromtimestamp(mtime, tz=timezone.utc).strftime("%Y-%m-%d")
        entries.append(
            "<url>"
            f"<loc>{SITE_URL}{page['path']}</loc>"
            f"<lastmod>{lastmod}</lastmod>"
            f"<priority>{page['priority']}</priority>"
            "</url>"
        )
    xml = (
        '<?xml version="1.0" encoding="UTF-8"?>'
        '<urlset xmlns="http://www.sitemaps.org/schemas/sitemap/0.9">'
        + "".join(entries)
        + "</urlset>"
    )
    return Response(xml, mimetype="application/xml")


@app.get("/blog/<slug>")
def blog_post(slug):
    post = BLOG_POSTS_BY_SLUG.get(slug)
    if not post:
        abort(404)
    return render_template(post["template"])


@app.get("/<slug>")
def keyword_page(slug):
    page = KEYWORD_PAGES_BY_SLUG.get(slug)
    if not page:
        abort(404)
    return render_template(page["template"], keyword=page["keyword"])


@app.errorhandler(413)
def request_too_large(_error):
    return jsonify(error="ファイルサイズは10MB以下にしてください。"), 413


@app.post("/api/ai-note")
def api_ai_note():
    try:
        prompt, image_part = get_ai_input_from_request()
    except ValueError as e:
        return jsonify(error=str(e)), 400

    try:
        client = create_gemini_client()
        response = client.models.generate_content(
            model=_GEMINI_MODEL,
            contents=build_ai_contents(_NOTE_PROMPT, prompt, image_part),
            config=genai_types.GenerateContentConfig(
                response_mime_type="application/json",
                max_output_tokens=4096,
                temperature=0.7,
            ),
        )
        raw = response.text.strip()
        if raw.startswith("```"):
            lines = raw.splitlines()
            raw = "\n".join(lines[1:-1] if lines[-1].startswith("```") else lines[1:])
        tree = json.loads(raw)
    except Exception as e:
        return jsonify(error=str(e)), 502

    return jsonify(tree=tree)


@app.post("/api/ai-mindmap")
def api_ai_mindmap():
    try:
        prompt, image_part = get_ai_input_from_request()
    except ValueError as e:
        return jsonify(error=str(e)), 400

    try:
        client = create_gemini_client()
        response = client.models.generate_content(
            model=_GEMINI_MODEL,
            contents=build_ai_contents(_MINDMAP_PROMPT, prompt, image_part),
            config=genai_types.GenerateContentConfig(
                response_mime_type="application/json",
                max_output_tokens=4096,
                temperature=0.7,
            ),
        )
        raw = response.text.strip()
        if raw.startswith("```"):
            lines = raw.splitlines()
            raw = "\n".join(lines[1:-1] if lines[-1].startswith("```") else lines[1:])
        tree = normalize_ai_mindmap_tree(json.loads(raw))
    except Exception as e:
        return jsonify(error=str(e)), 502

    return jsonify(tree=tree)


if __name__ == "__main__":
    port = int(os.getenv("PORT", "5006"))
    app.run(debug=True, port=port)
